"""
Task 06: DOCX to JSON Extractor
Extracts structured data from DOCX files and converts to JSON format.
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import argparse
import json
import re
from typing import Dict, List, Any, Optional
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from smolagents import ToolCallingAgent, InferenceClientModel


def extract_docx_content(file_path: str) -> Dict[str, Any]:
    """
    Extract content from a DOCX file.
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        Dictionary with extracted content
    """
    if not HAS_DOCX:
        raise ImportError("python-docx library required. Install with: pip install python-docx")
    
    doc = Document(file_path)
    
    # Extract paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "is_heading": para.style.name.startswith('Heading') if para.style else False
            })
    
    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    
    # Extract metadata
    metadata = {
        "file_name": Path(file_path).name,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "word_count": sum(len(p["text"].split()) for p in paragraphs)
    }
    
    return {
        "metadata": metadata,
        "paragraphs": paragraphs,
        "tables": tables
    }


def structure_with_ai(content: Dict[str, Any], extraction_type: str = "auto") -> Dict[str, Any]:
    """
    Use AI to structure the extracted content into meaningful JSON.
    
    Args:
        content: Extracted DOCX content
        extraction_type: Type of extraction (auto, resume, report, invoice, meeting)
    
    Returns:
        Structured JSON data
    """
    model = get_model()
    agent = ToolCallingAgent(tools=[], model=model, max_steps=3)
    
    # Prepare content preview
    text_content = "\n".join([p["text"] for p in content["paragraphs"][:20]])
    table_preview = json.dumps(content["tables"][:2]) if content["tables"] else "No tables"
    
    extraction_prompts = {
        "auto": "Analyze this document and extract structured data. Identify the document type and extract relevant fields.",
        "resume": "Extract resume/CV data including: name, contact info, skills, experience, education, certifications.",
        "report": "Extract report data including: title, executive_summary, key_findings, recommendations, sections.",
        "invoice": "Extract invoice data including: invoice_number, date, vendor, items (description, quantity, price), total.",
        "meeting": "Extract meeting notes including: date, attendees, agenda_items, action_items (with assignee, deadline), decisions."
    }
    
    prompt_instruction = extraction_prompts.get(extraction_type, extraction_prompts["auto"])
    
    prompt = f"""{prompt_instruction}

**Document Content:**
{text_content[:1500]}

**Tables:**
{table_preview[:500]}

Return ONLY valid JSON with extracted fields. Use appropriate nested structures for complex data like lists of items or sections."""
    
    try:
        result = agent.run(prompt)
        # Try to parse as JSON
        try:
            structured_data = json.loads(result)
        except json.JSONDecodeError:
            # Extract JSON from markdown code blocks if present
            json_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', result)
            if json_match:
                structured_data = json.loads(json_match.group(1))
            else:
                # Wrap raw text as structured data
                structured_data = {
                    "extracted_text": result,
                    "extraction_method": "ai_raw"
                }
    except Exception as e:
        structured_data = {
            "error": str(e),
            "raw_content": text_content[:500]
        }
    
    return {
        "source_file": content["metadata"]["file_name"],
        "extraction_type": extraction_type,
        "metadata": content["metadata"],
        "structured_data": structured_data
    }


def extract_docx_to_json(file_path: str, extraction_type: str = "auto", output_file: Optional[str] = None) -> str:
    """
    Extract structured JSON from a DOCX file.
    
    Args:
        file_path: Path to DOCX file
        extraction_type: Type of extraction (auto, resume, report, invoice, meeting)
        output_file: Optional path to save JSON output
    
    Returns:
        JSON string or error message
    """
    try:
        # Extract raw content
        content = extract_docx_content(file_path)
        
        # Structure with AI
        result = structure_with_ai(content, extraction_type)
        
        # Convert to JSON
        json_output = json.dumps(result, indent=2, ensure_ascii=False)
        
        # Save if requested
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(json_output)
        
        return json_output
        
    except ImportError as e:
        return f"Error: {e}"
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract structured JSON from DOCX files")
    parser.add_argument("file", help="Path to DOCX file")
    parser.add_argument("--type", choices=["auto", "resume", "report", "invoice", "meeting"],
                       default="auto", help="Extraction type")
    parser.add_argument("--output", help="Optional path to save JSON output")
    args = parser.parse_args()
    
    result = extract_docx_to_json(args.file, args.type, args.output)
    print(result)
