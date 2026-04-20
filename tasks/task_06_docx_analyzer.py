"""
Task 06: DOCX Analyzer
Extracts and analyzes content from Microsoft Word documents.
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import json
from docx import Document
from smolagents import ToolCallingAgent
from config import get_model


def run(file_path: str = "document.docx", extraction_type: str = "auto") -> str:
    """
    Extract and analyze content from a DOCX file.
    
    Args:
        file_path: Path to DOCX file
        extraction_type: Type of extraction (auto, resume, report, invoice, meeting)
        
    Returns:
        Extracted content and analysis
    """
    try:
        # Extract content from DOCX
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "is_heading": para.style.name.startswith('Heading') if para.style else False
                })
        
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        content = {
            "file_name": Path(file_path).name,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "word_count": sum(len(p["text"].split()) for p in paragraphs),
            "paragraphs": paragraphs[:20],  # Limit for AI processing
            "tables": tables[:2]
        }
        
        # Use AI to structure the content
        text_content = "\n".join([p["text"] for p in content["paragraphs"]])
        
        model = get_model()
        agent = ToolCallingAgent(tools=[], model=model, max_steps=3)
        
        prompts = {
            "auto": "Analyze this document and extract structured data. Identify document type and extract relevant fields.",
            "resume": "Extract resume data: name, contact, skills, experience, education.",
            "report": "Extract report data: title, summary, findings, recommendations.",
            "invoice": "Extract invoice data: number, date, vendor, items, total.",
            "meeting": "Extract meeting notes: date, attendees, agenda, action items, decisions."
        }
        
        prompt = f"""{prompts.get(extraction_type, prompts['auto'])}

Document content:
{text_content[:1500]}

Provide a structured analysis of this document."""
        
        analysis = agent.run(prompt)
        
        return f"""# DOCX Analysis Report

## Document Metadata
- **File:** {content['file_name']}
- **Paragraphs:** {content['paragraph_count']}
- **Tables:** {content['table_count']}
- **Word Count:** {content['word_count']}

## AI Analysis
{analysis}

## Content Preview
{text_content[:500]}...
"""
    except Exception as e:
        return f"Error analyzing DOCX: {str(e)}"


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="DOCX Analyzer")
    parser.add_argument("--file", default="document.docx", help="Path to DOCX file")
    parser.add_argument("--type", choices=["auto", "resume", "report", "invoice", "meeting"],
                       default="auto", help="Extraction type")
    args = parser.parse_args()
    print(run(args.file, args.type))
